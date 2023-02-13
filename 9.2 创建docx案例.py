# -*- coding: utf-8 -*-
"""
Created on Mon Feb 13 14:18:04 2023

@author: Mora

9.2 创建docx案例
"""

from docx import Document
from docx.shared import Cm

document = Document()

document.add_heading('富途女篮比赛结果', 0)

p = document.add_paragraph('2022年富途篮球队 ')
p.add_run('女生').bold = True
p.add_run(' 比赛 ')
p.add_run('结果公示').italic = True

document.add_heading('比赛信息', level=1)
document.add_paragraph('投篮比赛 & 3V3 & 带妹赛', style='Intense Quote')

document.add_paragraph(
    '无序列表的测试', style='List Bullet'
)
document.add_paragraph(
    '有序列表的测试', style='List Number'
)

document.add_picture('D:/富途证券/女生篮球赛/篮球.png', width=Cm(4.5))

document.add_section()

records = (
    ('1', '8', '小涵'),
    ('2', '12', '小敏'),
    ('3', '3', '小文')
)

table = document.add_table(rows=1, cols=3)
table.style = 'Colorful Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '编号'
hdr_cells[1].text = '进球数'
hdr_cells[2].text = '姓名'
for qty, id, name in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = name

document.add_page_break()

document.save('D:/富途证券/表格/demo.docx')


